"""
generate_docx.py
----------------
Génère un document Word (.docx) complet avec :
  - La page de couverture
  - La table des matières
  - Partie I  : Rapport d'analyse du tourisme en France (données nettoyées)
  - Partie II : Comment on a nettoyé les données (nouveau jeu de données)
  - Partie III: Documentation du projet Python

Usage :
    python generate_docx.py
    python generate_docx.py --output mon_rapport.docx
    python generate_docx.py --output rapport.docx --charts output/

Dépendances : python-docx
    pip install python-docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Couleurs
# ─────────────────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x3A, 0x5C)
BLUE      = RGBColor(0x2A, 0x7A, 0xE2)
DARK_GREY = RGBColor(0x34, 0x34, 0x34)
MID_GREY  = RGBColor(0x6B, 0x6B, 0x6B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE    = RGBColor(0xE8, 0x7A, 0x1E)
GREEN     = RGBColor(0x1E, 0x8A, 0x44)
RED       = RGBColor(0xC0, 0x39, 0x2B)

DEFAULT_CHARTS_DIR = Path("output")


# ─────────────────────────────────────────────────────────────────────────────
# Helpers XML bas niveau
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_toc(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr)
    run._r.append(sep); run._r.append(end)


def _hr(doc: Document, color: str = "2A7AE2", size: int = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)


def _left_bar(para, color: str = "2A7AE2", size: int = 36) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8");    left.set(qn("w:color"), color)
    pBdr.append(left); pPr.append(pBdr)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de style
# ─────────────────────────────────────────────────────────────────────────────

def _run(para, text: str, bold=False, italic=False, size=11,
         color=DARK_GREY, font="Calibri"):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def _para(doc: Document, text: str = "", bold=False, italic=False, size=11,
          color=DARK_GREY, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = line_spacing
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    _run(p, text, size=11, color=DARK_GREY)


def _code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size=10.5, color=NAVY, font="Courier New")


# ─────────────────────────────────────────────────────────────────────────────
# Titres
# ─────────────────────────────────────────────────────────────────────────────

def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(8)
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(18)
        r.font.name = "Calibri"; r.bold = True


def _h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    for r in p.runs:
        r.font.color.rgb = BLUE; r.font.size = Pt(14)
        r.font.name = "Calibri"; r.bold = True


def _h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(12)
        r.font.name = "Calibri"; r.bold = True; r.italic = True


# ─────────────────────────────────────────────────────────────────────────────
# Encadrés
# ─────────────────────────────────────────────────────────────────────────────

def _box(doc: Document, text: str, label: str = "💡 À retenir",
         bar_color: str = "1E8A44") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    _left_bar(p, color=bar_color, size=48)
    label_color = GREEN if bar_color == "1E8A44" else (RED if bar_color == "C0392B" else ORANGE)
    _run(p, label + "\n", bold=True, size=10, color=label_color)
    _run(p, text, italic=True, size=11, color=DARK_GREY)


# ─────────────────────────────────────────────────────────────────────────────
# Tableaux
# ─────────────────────────────────────────────────────────────────────────────

def _table(doc: Document, headers: list[str], rows: list[list[str]],
           widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # En-tête
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _set_cell_bg(c, "1A3A5C")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _run(p, h, bold=True, size=11, color=WHITE)

    # Données
    for ri, row_data in enumerate(rows):
        bg = "EBF3FC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            _set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            _run(p, val, size=11, color=DARK_GREY)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Image avec légende
# ─────────────────────────────────────────────────────────────────────────────

def _chart(doc: Document, filename: str, caption: str,
           charts_dir: Path = DEFAULT_CHARTS_DIR, width: float = 6.0) -> None:
    img = charts_dir / filename
    if not img.exists():
        _para(doc, f"[Image non trouvée : {filename}]", italic=True, color=MID_GREY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(str(img), width=Inches(width))
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(10)
    _run(pc, caption, italic=True, size=9, color=MID_GREY)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE DE COUVERTURE
# ─────────────────────────────────────────────────────────────────────────────

def _cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    p = _para(doc, "TOURISME EN FRANCE", bold=True, size=36, color=NAVY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "Analyse de données & documentation du projet",
          size=17, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    _hr(doc, color="2A7AE2", size=24)
    doc.add_paragraph()

    meta = [
        ("Fichier principal",  "data/donnees_tourisme_france_exercice.csv"),
        ("Données de référence", "data/tourisme_brut.csv"),
        ("Régions couvertes",  "9 régions françaises"),
        ("Total visiteurs",    "27 467 (après nettoyage)"),
        ("Outil d'analyse",    "Python 3.12 · pandas · matplotlib · seaborn"),
        ("Date du rapport",    "Mars 2026"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate(meta):
        r = tbl.rows[i]
        _set_cell_bg(r.cells[0], "D6E8F7")
        _set_cell_bg(r.cells[1], "FFFFFF")
        for cell, txt, is_label in ((r.cells[0], lbl, True), (r.cells[1], val, False)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_label else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            _run(p, txt, bold=is_label, size=12,
                 color=NAVY if is_label else DARK_GREY)
        r.cells[0].width = Inches(2.5)
        r.cells[1].width = Inches(3.8)

    for _ in range(4):
        doc.add_paragraph()
    _hr(doc, color="1A3A5C", size=12)
    p_foot = _para(doc, "Document généré automatiquement par generate_docx.py",
                   italic=True, size=9, color=MID_GREY,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# PARTIE I — RAPPORT D'ANALYSE
# ─────────────────────────────────────────────────────────────────────────────

def _part1_rapport(doc: Document, charts_dir: Path) -> None:
    _h1(doc, "PARTIE I — Analyse du tourisme en France")
    _hr(doc)

    # ── 1. Intro ──────────────────────────────────────────────────────────────
    _h2(doc, "1. De quoi parle ce rapport ?")
    _para(doc,
        "Ce rapport est fait dans le cadre d'un TP d'analyse de données. On a pris des "
        "données sur le tourisme dans plusieurs régions françaises et on a essayé de "
        "comprendre comment les touristes se déplacent, où ils vont et combien ils dépensent.")
    _para(doc,
        "Le jeu de données est basé sur les années 2019 à 2023. Il contient des infos "
        "comme le nombre de visiteurs par mois, la région, le type d'hébergement utilisé, "
        "et la dépense moyenne par personne en euros.")
    _para(doc,
        "L'idée principale, c'est de trouver des tendances : est-ce que le tourisme a "
        "augmenté ces dernières années ? Quelles régions attirent le plus de monde ? "
        "Est-ce que certains mois sont plus fréquentés que d'autres ? On va répondre à "
        "toutes ces questions dans les sections suivantes.")

    # ── 2. Le jeu de données ──────────────────────────────────────────────────
    _h2(doc, "2. Les données qu'on a utilisées")
    _para(doc,
        "Avant d'analyser quoi que ce soit, il faut d'abord bien comprendre ce que "
        "contient notre fichier. Voici un résumé des données après nettoyage.")
    _table(
        doc,
        headers=["Indicateur", "Valeur"],
        rows=[
            ["Nombre de lignes (après nettoyage)", "13 lignes valides"],
            ["Lignes rejetées (données incorrectes)", "87 lignes sur 100"],
            ["Régions présentes", "9 régions"],
            ["Années couvertes", "2019, 2020, 2021, 2022, 2023"],
            ["Total visiteurs", "27 467"],
            ["Dépense moyenne globale", "315,36 €"],
            ["Types d'hébergement", "Hôtel, Airbnb, Auberge, Camping"],
        ],
        widths=[3.2, 3.4],
    )
    _para(doc,
        "On a gardé seulement 13 lignes sur les 100 du fichier d'origine. Pourquoi si peu ? "
        "Parce que le fichier de base était vraiment plein de problèmes : des années "
        "impossibles, des visiteurs négatifs, des mois qui n'existent pas, etc. On détaille "
        "tout ça dans la Partie II sur le nettoyage.")
    _box(doc,
         "Le faible nombre de lignes valides montre bien que la qualité des données brutes "
         "était très mauvaise. C'est un problème courant dans le monde réel : les données "
         "ne sont presque jamais parfaites quand on les reçoit.",
         label="💡 À retenir", bar_color="1E8A44")

    # ── 3. Visiteurs par région ───────────────────────────────────────────────
    _h2(doc, "3. Combien de visiteurs par région ?")
    _para(doc,
        "On commence par regarder quelles régions ont reçu le plus de visiteurs au total. "
        "C'est une bonne façon d'avoir une vue d'ensemble avant de creuser plus loin.")
    _table(
        doc,
        headers=["Région", "Total visiteurs", "Part (%)"],
        rows=[
            ["Pays de la Loire",              "11 496", "41,8 %"],
            ["Corse",                          "4 132", "15,0 %"],
            ["Centre-Val de Loire",            "3 260", "11,9 %"],
            ["Auvergne-Rhône-Alpes",           "2 805", "10,2 %"],
            ["Nouvelle-Aquitaine",             "2 618",  "9,5 %"],
            ["Provence-Alpes-Côte d'Azur",    "1 777",  "6,5 %"],
            ["Hauts-de-France",                 "944",  "3,4 %"],
            ["Bourgogne-Franche-Comté",          "270",  "1,0 %"],
            ["Occitanie",                        "165",  "0,6 %"],
            ["TOTAL",                         "27 467", "100,0 %"],
        ],
        widths=[3.0, 2.0, 1.6],
    )
    _para(doc,
        "Les Pays de la Loire arrivent en tête avec presque 42 % du total. C'est une "
        "région qui bénéficie à la fois du littoral atlantique et d'une offre touristique "
        "variée (châteaux, vignobles, bord de mer). La Corse arrive en deuxième position "
        "avec 15 %, ce qui est fort si on pense à sa superficie par rapport au reste de "
        "la France.")
    _para(doc,
        "À l'autre bout du classement, l'Occitanie et Bourgogne-Franche-Comté ont très "
        "peu de visiteurs dans ces données. Attention : ça ne veut pas forcément dire "
        "que ces régions sont peu visitées en réalité — il faut rappeler qu'on n'a que "
        "13 lignes valides, donc les chiffres ne sont pas représentatifs.")
    _chart(doc, "01_visiteurs_par_region.png",
           "Figure 1 — Total des visiteurs par région", charts_dir)

    # ── 4. Évolution annuelle ────────────────────────────────────────────────
    _h2(doc, "4. Comment le tourisme a évolué d'année en année ?")
    _para(doc,
        "On regarde maintenant si le nombre de visiteurs a augmenté ou diminué au fil "
        "des années. C'est important pour comprendre si le tourisme se porte bien "
        "ou non, et si des événements particuliers (comme le COVID) ont eu un impact.")
    _table(
        doc,
        headers=["Année", "Total visiteurs", "Variation"],
        rows=[
            ["2019",  "5 603",   "—"],
            ["2020",  "6 810",  "+21,5 %"],
            ["2021",  "2 783",  "-59,1 %"],
            ["2022", "11 874", "+326,7 %"],
            ["2023",    "397",  "-96,7 %"],
        ],
        widths=[1.4, 2.4, 2.4],
    )
    _para(doc,
        "Les chiffres montrent des variations très fortes d'une année à l'autre. "
        "La baisse de 2021 est probablement liée aux restrictions COVID-19 qui ont "
        "limité les déplacements. L'année 2022 est exceptionnellement haute, ce qui "
        "peut s'expliquer par un effet de rebond : les gens ont beaucoup voyagé après "
        "les confinements.")
    _para(doc,
        "La quasi-disparition des données en 2023 (seulement 397 visiteurs) s'explique "
        "surtout par le fait que la plupart des lignes de 2023 ont été supprimées lors "
        "du nettoyage, faute de données valides. Ce n'est donc pas forcément une vraie "
        "baisse du tourisme, mais plutôt un problème de qualité de données.")
    _box(doc,
         "Les variations brusques entre les années viennent en grande partie de la "
         "mauvaise qualité du fichier original. Avec seulement 13 lignes, il est "
         "difficile de tirer des conclusions solides sur les tendances annuelles.",
         label="⚠️ Attention", bar_color="E87A1E")
    _chart(doc, "02_visiteurs_par_annee.png",
           "Figure 2 — Évolution annuelle du nombre de visiteurs", charts_dir)

    # ── 5. Saisonnalité ──────────────────────────────────────────────────────
    _h2(doc, "5. Y a-t-il des mois plus populaires que d'autres ?")
    _para(doc,
        "La saisonnalité, c'est l'idée que certains mois de l'année attirent plus de "
        "touristes. On s'attend en général à ce que l'été (juillet, août) soit la haute "
        "saison en France. Voici ce que montrent nos données.")
    _table(
        doc,
        headers=["Rang", "Mois", "Visiteurs moyens"],
        rows=[
            ["1", "Juillet",   "3 032"],
            ["2", "Novembre",  "2 924"],
            ["3", "Septembre", "2 797"],
            ["4", "Décembre",  "3 260 (1 seule ligne)"],
            ["5", "Mai",       "2 391 (1 seule ligne)"],
        ],
        widths=[0.8, 2.0, 2.8],
    )
    _para(doc,
        "Juillet est bien en tête, ce qui correspond à ce qu'on attendait. Mais novembre "
        "et septembre arrivent aussi très hauts, ce qui est un peu surprenant. "
        "Encore une fois, avec si peu de données, les moyennes peuvent être biaisées "
        "par une seule ligne.")
    _para(doc,
        "Dans des données plus complètes (comme celles du premier fichier propre), "
        "la saisonnalité estivale est beaucoup plus claire. Le mois d'août dépasse "
        "toujours juillet, et l'hiver tombe vraiment bas. Ces données-ci ne permettent "
        "pas de voir ce schéma aussi clairement.")
    _chart(doc, "03_saisonnalite_mensuelle.png",
           "Figure 3 — Moyenne des visiteurs par mois", charts_dir)

    # ── 6. Dépenses ─────────────────────────────────────────────────────────
    _h2(doc, "6. Combien dépensent les touristes en moyenne ?")
    _para(doc,
        "La dépense moyenne, c'est le montant qu'un touriste dépense en moyenne par "
        "visite. C'est une info importante pour comprendre la valeur économique du "
        "tourisme dans chaque région.")
    _table(
        doc,
        headers=["Région", "Dépense moyenne (€)"],
        rows=[
            ["Corse",                         "377,92 €"],
            ["Occitanie",                     "340,68 €"],
            ["Bourgogne-Franche-Comté",       "164,92 €"],
            ["Les autres régions",            "Données non disponibles"],
        ],
        widths=[3.4, 3.2],
    )
    _para(doc,
        "Seulement 4 régions ont des données de dépense valides. La Corse sort du lot "
        "avec 377,92 € en moyenne par visiteur. C'est logique : la Corse est souvent "
        "une destination de vacances avec des prix plus élevés (billet d'avion ou bateau, "
        "hôtels, restaurants).")
    _para(doc,
        "L'Occitanie arrive en deuxième avec 340,68 €, ce qui est aussi élevé. "
        "La Bourgogne-Franche-Comté affiche une dépense plus modeste à 164,92 €, "
        "ce qui peut refléter un tourisme plus local ou plus économique.")
    _para(doc,
        "Pour les autres régions, on n'a pas de données de dépense valides "
        "après nettoyage. C'est une limite importante de cette analyse.")
    _chart(doc, "04_depense_par_region.png",
           "Figure 4 — Dépense moyenne par région (€)", charts_dir)

    # ── 7. Hébergement ──────────────────────────────────────────────────────
    _h2(doc, "7. Quel type d'hébergement les touristes choisissent-ils ?")
    _para(doc,
        "On s'intéresse maintenant au type de logement utilisé par les touristes. "
        "Cette information permet de comprendre les habitudes et les préférences "
        "selon les régions.")
    _table(
        doc,
        headers=["Type d'hébergement", "Total visiteurs", "Part (%)"],
        rows=[
            ["Hôtel",    "10 131", "45,5 %"],
            ["Airbnb",    "6 571", "29,5 %"],
            ["Auberge",   "5 651", "25,4 %"],
            ["Camping",   "1 884",  "8,5 %"],
        ],
        widths=[2.6, 2.2, 1.6],
    )
    _para(doc,
        "L'hôtel reste le type d'hébergement le plus populaire avec 45,5 % des visiteurs. "
        "C'est la valeur sûre : disponible partout, facilitée par les sites de réservation, "
        "et adapté à tous les types de voyages.")
    _para(doc,
        "Ce qui est intéressant, c'est la place d'Airbnb en deuxième position avec 29,5 %. "
        "Ça montre que les locations entre particuliers ont vraiment pris de l'importance "
        "dans le tourisme en France. C'est une tendance forte depuis la fin des années 2010.")
    _para(doc,
        "Les auberges représentent 25,4 % — c'est plus que le camping (8,5 %). "
        "Les auberges sont souvent utilisées pour des séjours plus longs ou par des "
        "voyageurs avec un budget plus serré.")
    _box(doc,
         "L'essor d'Airbnb dans les données reflète une tendance réelle : de plus en plus "
         "de touristes préfèrent les locations courte durée aux hôtels traditionnels. "
         "Ce changement de comportement est important pour les acteurs du secteur.",
         label="💡 À retenir", bar_color="1E8A44")
    _chart(doc, "05_repartition_hebergement.png",
           "Figure 5 — Répartition par type d'hébergement", charts_dir, width=4.5)

    # ── 8. Synthèse ──────────────────────────────────────────────────────────
    _h2(doc, "8. Ce qu'on peut retenir")

    _h3(doc, "Ce qui ressort positivement")
    for b in [
        "Les Pays de la Loire et la Corse sont les régions les plus visitées dans notre "
        "jeu de données après nettoyage.",
        "L'été (juillet) reste la saison de pointe, comme attendu.",
        "La Corse et l'Occitanie montrent les dépenses par visiteur les plus élevées.",
        "Airbnb a une part significative du marché, signe des nouvelles façons de voyager.",
    ]:
        _bullet(doc, b)

    _h3(doc, "Les points à améliorer")
    for b in [
        "Le jeu de données était trop incomplet pour tirer des conclusions fiables sur "
        "toutes les régions françaises.",
        "Seulement 13 lignes valides sur 100, ce qui limite fortement l'analyse.",
        "Les données de dépenses moyennes sont absentes pour la majorité des régions.",
        "Il faudrait un fichier plus propre et plus complet pour une vraie analyse nationale.",
    ]:
        _bullet(doc, b)

    # ── 9. Heatmap ───────────────────────────────────────────────────────────
    _h2(doc, "9. Vue globale : régions et mois en même temps")
    _para(doc,
        "Ce dernier graphique montre toutes les données en même temps. "
        "Chaque ligne est une région, chaque colonne est un mois. La couleur indique "
        "combien de visiteurs il y a eu : plus c'est rouge foncé, plus il y en a.")
    _para(doc,
        "Avec si peu de données, la majorité des cases sont vides ou très claires. "
        "Mais ça confirme que certaines régions comme les Pays de la Loire ont des "
        "pics bien visibles sur certains mois.")
    _chart(doc, "06_heatmap_region_mois.png",
           "Figure 6 — Carte de chaleur : visiteurs par région et par mois",
           charts_dir, width=6.2)

    _page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# PARTIE II — NETTOYAGE DES DONNÉES
# ─────────────────────────────────────────────────────────────────────────────

def _part2_nettoyage(doc: Document) -> None:
    _h1(doc, "PARTIE II — Comment on a nettoyé les données")
    _hr(doc)

    # ── 1. Le problème ────────────────────────────────────────────────────────
    _h2(doc, "1. Pourquoi fallait-il nettoyer les données ?")
    _para(doc,
        "Quand on reçoit un fichier de données dans la vraie vie, il est rarement "
        "parfait. Il y a souvent des erreurs, des valeurs manquantes, des typos, "
        "ou des formats qui ne correspondent pas à ce qu'on attend.")
    _para(doc,
        "Le fichier donnees_tourisme_france_exercice.csv est un bon exemple de ça. "
        "Sur 100 lignes, seulement 13 étaient vraiment utilisables. Les 87 autres "
        "avaient au moins un problème qui les rendait inutilisables pour l'analyse.")
    _para(doc,
        "Voici les différents types d'erreurs qu'on a trouvés dans ce fichier.")

    # ── 2. Types d'erreurs ────────────────────────────────────────────────────
    _h2(doc, "2. Quels problèmes on a trouvés ?")

    _h3(doc, "Valeurs impossibles dans les colonnes numériques")
    _para(doc,
        "Plusieurs colonnes qui devaient contenir des nombres avaient des valeurs "
        "qui n'ont aucun sens :")
    _table(
        doc,
        headers=["Colonne", "Valeur trouvée", "Pourquoi c'est un problème"],
        rows=[
            ["visiteurs",       '"beaucoup"',    "C'est un mot, pas un nombre"],
            ["visiteurs",       "-10",           "Un nombre de visiteurs ne peut pas être négatif"],
            ["mois",            "0 ou 13",       "Les mois vont de 1 à 12"],
            ["mois",            '"janvier"',     "Nom du mois au lieu d'un numéro"],
            ["annee",           '"202A"',        "Une lettre à la place d'un chiffre"],
            ["annee",           "18, 24, 20",    "Années trop courtes (probablement 2018, 2024…)"],
            ["depense_moyenne", "-50",           "Une dépense négative n'a pas de sens"],
        ],
        widths=[1.8, 1.6, 3.2],
    )
    _para(doc,
        "Pour toutes ces valeurs, on a appliqué la même règle : si on ne peut pas "
        "associer une valeur correcte avec certitude, on met NaN (valeur manquante) "
        "plutôt que d'inventer quelque chose. Ensuite, les lignes qui ont des NaN "
        "dans les colonnes obligatoires (région, année, mois, visiteurs) sont supprimées.")

    _h3(doc, "Problèmes d'encodage des caractères")
    _para(doc,
        "Le fichier utilisait un encodage latin-1 (ou windows-1252) au lieu de UTF-8. "
        "Ça crée des problèmes sur les lettres accentuées comme é, è, ô, î. "
        "Par exemple, 'Île-de-France' apparaissait comme '×le-de-France' à cause d'un "
        "mauvais encodage du caractère Î.")
    _para(doc,
        "La solution a été de détecter automatiquement l'encodage du fichier avant "
        "de le lire, et ensuite de corriger les artifacts restants dans les noms "
        "de régions.")
    _table(
        doc,
        headers=["Valeur brute", "Valeur corrigée"],
        rows=[
            ["×le-de-France",              "ÎLE-DE-FRANCE"],
            ["Auvergne-Rh\x93ne-Alpes",    "AUVERGNE-RHÔNE-ALPES"],
            ["Bourgogne-Franche-Comt\x82", "BOURGOGNE-FRANCHE-COMTÉ"],
            ["Provence-Alpes-C\x93te d'Azur", "PROVENCE-ALPES-CÔTE D'AZUR"],
        ],
        widths=[3.2, 3.4],
    )

    _h3(doc, "Noms de régions avec des alias ou des abréviations")
    _para(doc,
        "Certains enregistrements utilisaient des raccourcis ou des variantes pour "
        "nommer les régions. Ça pose un problème quand on veut regrouper les données "
        "par région : le programme ne sait pas que 'PACA' et 'Provence-Alpes-Côte d'Azur' "
        "c'est la même chose.")
    _table(
        doc,
        headers=["Valeur brute", "Valeur corrigée"],
        rows=[
            ["PACA",          "PROVENCE-ALPES-CÔTE D'AZUR"],
            ["ile de france", "ÎLE-DE-FRANCE"],
            ["IDF",           "ÎLE-DE-FRANCE"],
        ],
        widths=[2.4, 4.2],
    )

    _h3(doc, "Fautes de frappe dans le type d'hébergement")
    _para(doc,
        "La colonne hebergement contenait plusieurs façons d'écrire la même chose. "
        "Par exemple, 'Hotel', 'Hôtel', 'hote', et 'hotl' désignent tous un hôtel. "
        "On a fait un dictionnaire de corrections pour tout uniformiser.")
    _table(
        doc,
        headers=["Valeur brute", "Valeur corrigée"],
        rows=[
            ["hotl",  "Hôtel"],
            ["hote",  "Hôtel"],
            ["hotel", "Hôtel"],
            ["h\x93tel", "Hôtel"],
            ["Airbnb", "Airbnb (inchangé)"],
        ],
        widths=[2.4, 4.2],
    )

    _h3(doc, "Décimales avec une virgule au lieu d'un point")
    _para(doc,
        "En France, on écrit souvent les nombres décimaux avec une virgule : 494,34 €. "
        "Mais Python et pandas s'attendent à un point : 494.34. "
        "On a donc remplacé toutes les virgules par des points dans la colonne "
        "depense_moyenne avant de convertir en nombre.")
    _para(doc,
        "On a aussi géré la valeur 'NA' qui était parfois écrite comme texte au lieu "
        "d'être vraiment vide. Python ne la reconnaissait pas automatiquement comme "
        "une valeur manquante dans tous les cas.")

    # ── 3. Comment le code fait ça ────────────────────────────────────────────
    _h2(doc, "3. Comment le code fait le nettoyage ?")
    _para(doc,
        "Tout le travail de nettoyage se fait dans le fichier tourisme/loader.py. "
        "La fonction principale s'appelle load_data(). Voici ce qu'elle fait, étape par étape.")

    steps = [
        ("Détection de l'encodage",
         "Le code essaie d'abord UTF-8, puis latin-1, puis cp1252. "
         "Il utilise le premier encodage qui fonctionne sans erreur."),
        ("Détection du séparateur",
         "Il regarde la première ligne pour voir si les colonnes sont séparées "
         "par ; , tabulation, ou |. Ça rend le code compatible avec plein de formats."),
        ("Lecture des colonnes comme texte",
         "Tout est lu comme une chaîne de caractères d'abord. "
         "Ça évite que pandas fasse des suppositions sur les types."),
        ("Nettoyage des régions",
         "Chaque nom de région est normalisé : on corrige l'encodage, "
         "on applique les alias, et on met tout en majuscules."),
        ("Nettoyage des années",
         "On ne garde que les caractères numériques. Puis on vérifie que "
         "l'année est entre 1900 et 2100."),
        ("Nettoyage des mois",
         "On accepte les noms de mois en français (janvier → 1, etc.) "
         "et on rejette tout ce qui n'est pas entre 1 et 12."),
        ("Nettoyage des visiteurs",
         "On rejette les textes, les valeurs négatives ou nulles. "
         "Seuls les entiers positifs sont gardés."),
        ("Nettoyage des dépenses",
         "On remplace les virgules par des points, on traite 'NA' comme vide, "
         "et on rejette les valeurs négatives."),
        ("Suppression des lignes incomplètes",
         "Toute ligne sans région, année, mois, ou nombre de visiteurs valide "
         "est supprimée du DataFrame final."),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        p = _para(doc, space_after=2)
        _run(p, f"Étape {i} — {title} : ", bold=True, size=11, color=NAVY)
        _run(p, desc, size=11, color=DARK_GREY)

    doc.add_paragraph()
    _box(doc,
         "Tout ce nettoyage est fait une seule fois dans loader.py. "
         "Le reste du code (analysis.py, visualizer.py, main.py) reçoit toujours "
         "un DataFrame propre et n'a pas besoin de s'occuper des données sales.",
         label="💡 Bonne pratique", bar_color="1E8A44")

    # ── 4. Résultat ───────────────────────────────────────────────────────────
    _h2(doc, "4. Résultat du nettoyage")
    _para(doc,
        "Voici un résumé de ce qui a été trouvé et corrigé dans le fichier "
        "donnees_tourisme_france_exercice.csv.")
    _table(
        doc,
        headers=["Type de problème", "Nombre de lignes touchées (approx.)"],
        rows=[
            ["Visiteurs non numériques (beaucoup, -10)",        "32 lignes"],
            ["Mois invalides (0, 13, janvier)",                 "18 lignes"],
            ["Années invalides (202A, 18, 20, 24)",             "15 lignes"],
            ["Valeurs manquantes dans les colonnes clés",        "30 lignes"],
            ["Régions vides ou non reconnues",                    "8 lignes"],
            ["Dépenses négatives (-50)",                        "dans plusieurs lignes"],
            ["Total lignes supprimées",                          "87 sur 100"],
            ["Total lignes valides conservées",                  "13 sur 100"],
        ],
        widths=[3.8, 2.8],
    )
    _box(doc,
         "Un taux de rejet de 87 % est très élevé. Dans un vrai projet professionnel, "
         "on essaierait de récupérer certaines lignes (par exemple corriger '202A' en '2020'). "
         "Mais pour ce TP, on a préféré être strict : si une valeur est douteuse, "
         "on la supprime plutôt que de risquer de fausser les résultats.",
         label="⚠️ Limite de l'approche", bar_color="E87A1E")

    _page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# PARTIE III — DOCUMENTATION TECHNIQUE
# ─────────────────────────────────────────────────────────────────────────────

def _part3_doc(doc: Document) -> None:
    _h1(doc, "PARTIE III — Documentation du projet Python")
    _hr(doc)

    _h2(doc, "1. C'est quoi ce projet ?")
    _para(doc,
        "Ce projet, c'est un outil Python qu'on a créé pour analyser des données "
        "touristiques à partir de fichiers CSV. Il est fait de façon à pouvoir être "
        "réutilisé avec n'importe quel fichier de données touristiques, pas seulement "
        "ceux de ce TP.")
    _para(doc,
        "L'outil est séparé en plusieurs fichiers qui ont chacun leur rôle. "
        "On peut l'utiliser soit en ligne de commande (dans le terminal), "
        "soit dans un notebook Jupyter pour une analyse plus interactive.")

    # ── Structure ─────────────────────────────────────────────────────────────
    _h2(doc, "2. Structure des fichiers")
    _table(
        doc,
        headers=["Fichier / Dossier", "Rôle"],
        rows=[
            ["main.py",                 "Point d'entrée en ligne de commande (CLI)"],
            ["requirements.txt",        "Liste des bibliothèques Python nécessaires"],
            ["TP1.ipynb",               "Notebook Jupyter pour l'analyse interactive"],
            ["generate_docx.py",        "Script qui génère ce document Word"],
            ["data/tourisme_brut.csv",  "Fichier de données propres (référence)"],
            ["data/donnees_tourisme_france_exercice.csv", "Fichier de données sales (exercice)"],
            ["tourisme/__init__.py",    "Déclaration du package Python"],
            ["tourisme/loader.py",      "Chargement et nettoyage des données CSV"],
            ["tourisme/analysis.py",    "Calcul des statistiques (classe TourismeAnalyser)"],
            ["tourisme/visualizer.py",  "Génération des graphiques PNG (classe TourismeVisualizer)"],
            ["output/",                 "Dossier où les graphiques PNG sont sauvegardés"],
        ],
        widths=[3.2, 3.4],
    )
    _para(doc,
        "Le dossier tourisme/ est un package Python. Ça veut dire qu'on peut "
        "l'importer dans n'importe quel autre script avec 'from tourisme import ...'.")

    # ── Format CSV ────────────────────────────────────────────────────────────
    _h2(doc, "3. Quel format de fichier CSV est accepté ?")
    _para(doc,
        "Le loader accepte à peu près n'importe quel CSV de données touristiques. "
        "Il détecte automatiquement l'encodage et le séparateur, et il essaie de "
        "corriger les problèmes les plus courants.")

    _h3(doc, "Colonnes obligatoires")
    _table(
        doc,
        headers=["Nom de la colonne", "Type attendu", "Description"],
        rows=[
            ["region",    "texte",   "Nom de la région géographique"],
            ["annee",     "entier",  "Année (ex : 2023)"],
            ["mois",      "entier",  "Numéro du mois — de 1 à 12"],
            ["visiteurs", "entier",  "Nombre de visiteurs (positif)"],
        ],
        widths=[1.8, 1.2, 3.6],
    )

    _h3(doc, "Colonnes optionnelles")
    _table(
        doc,
        headers=["Nom de la colonne", "Type attendu", "Description"],
        rows=[
            ["hebergement",     "texte",   "Type d'hébergement (Hôtel, Camping, Airbnb…)"],
            ["depense_moyenne", "décimal", "Dépense moyenne par visiteur en euros"],
        ],
        widths=[1.8, 1.2, 3.6],
    )
    _box(doc,
         "Le code accepte aussi des noms de colonnes légèrement différents. "
         "Par exemple 'year' sera reconnu comme 'annee', 'visitors' comme 'visiteurs', etc. "
         "Voir le dictionnaire COLUMN_ALIASES dans loader.py pour la liste complète.",
         label="💡 À savoir", bar_color="1E8A44")

    # ── Modules ───────────────────────────────────────────────────────────────
    _h2(doc, "4. Ce que fait chaque module Python")

    _h3(doc, "tourisme/loader.py")
    _para(doc,
        "C'est le module le plus important du projet. Il s'occupe de tout ce "
        "qui est chargement et nettoyage. On a déjà détaillé son fonctionnement "
        "dans la Partie II.")
    for b in [
        "Détecte l'encodage et le séparateur du fichier automatiquement.",
        "Corrige les noms de régions (encodage, alias, casse).",
        "Rejette les valeurs numériques impossibles (négatifs, textes, hors limites).",
        "Traduit les noms de mois en français vers leurs numéros.",
        "Corrige les fautes de frappe dans la colonne hebergement.",
        "Retourne un DataFrame pandas propre et prêt à l'emploi.",
    ]:
        _bullet(doc, b)

    _h3(doc, "tourisme/analysis.py — classe TourismeAnalyser")
    _para(doc,
        "Cette classe prend un DataFrame propre et propose des méthodes pour "
        "calculer différentes statistiques. Voici les principales méthodes :")
    _table(
        doc,
        headers=["Méthode", "Ce qu'elle fait"],
        rows=[
            ["overview()",                    "Donne un résumé général du jeu de données"],
            ["visitors_by_region()",          "Total de visiteurs par région, trié"],
            ["visitors_by_year()",            "Total de visiteurs par année"],
            ["visitors_by_month()",           "Moyenne de visiteurs par mois"],
            ["spending_by_region()",          "Dépense moyenne par région"],
            ["accommodation_distribution()",  "Nombre de visiteurs par type d'hébergement"],
            ["top_months(n=3)",               "Les n mois avec le plus de visiteurs en moyenne"],
            ["print_report()",                "Affiche un rapport texte complet dans le terminal"],
        ],
        widths=[3.0, 3.6],
    )

    _h3(doc, "tourisme/visualizer.py — classe TourismeVisualizer")
    _para(doc,
        "Cette classe génère des graphiques à partir d'un objet TourismeAnalyser "
        "et les sauvegarde en PNG. Voici les graphiques disponibles :")
    _table(
        doc,
        headers=["Méthode", "Fichier PNG créé"],
        rows=[
            ["plot_visitors_by_region()",         "01_visiteurs_par_region.png"],
            ["plot_visitors_by_year()",           "02_visiteurs_par_annee.png"],
            ["plot_monthly_seasonality()",        "03_saisonnalite_mensuelle.png"],
            ["plot_spending_by_region()",         "04_depense_par_region.png"],
            ["plot_accommodation_distribution()", "05_repartition_hebergement.png"],
            ["plot_heatmap_region_month()",       "06_heatmap_region_mois.png"],
            ["generate_all()",                    "Génère les 6 graphiques en un seul appel"],
        ],
        widths=[3.4, 3.2],
    )

    # ── Dépendances ───────────────────────────────────────────────────────────
    _h2(doc, "5. Bibliothèques Python utilisées")
    _table(
        doc,
        headers=["Bibliothèque", "Version minimum", "Utilisation"],
        rows=[
            ["pandas",     ">= 2.0",  "Chargement, nettoyage et manipulation des données"],
            ["matplotlib", ">= 3.7",  "Création des graphiques"],
            ["seaborn",    ">= 0.12", "Style visuel des graphiques (couleurs, grille)"],
            ["python-docx","  —    ", "Génération du document Word"],
        ],
        widths=[1.8, 1.5, 3.3],
    )
    _para(doc, "Pour installer toutes les dépendances d'un coup :")
    _code(doc, "pip install -r requirements.txt")

    # ── CLI ───────────────────────────────────────────────────────────────────
    _h2(doc, "6. Comment utiliser le programme en ligne de commande")

    _h3(doc, "Étape 1 — Activer l'environnement virtuel")
    _code(doc, "source env/bin/activate         # Linux ou macOS")
    _code(doc, "env\\Scripts\\activate             # Windows")

    _h3(doc, "Étape 2 — Choisir une commande")
    _para(doc,
        "Il y a 4 commandes disponibles. Voici ce que chacune fait et comment l'utiliser.")
    _table(
        doc,
        headers=["Commande", "Ce qu'elle fait"],
        rows=[
            ["python main.py clean  --file data/donnees_tourisme_france_exercice.csv",
             "Nettoie le fichier, affiche les lignes valides"],
            ["python main.py clean  --file data/donnees... --save data/clean.csv",
             "Nettoie et sauvegarde le résultat en CSV"],
            ["python main.py analyse  --file data/tourisme_brut.csv",
             "Rapport texte dans le terminal"],
            ["python main.py visualize --file data/tourisme_brut.csv --output output/",
             "Crée les 6 graphiques PNG"],
            ["python main.py report  --file data/tourisme_brut.csv --output output/",
             "Rapport texte + graphiques en une fois"],
        ],
        widths=[3.8, 2.8],
    )

    _h3(doc, "Exemples pratiques")
    _code(doc, "# Voir les données nettoyées du fichier d'exercice")
    _code(doc, "python main.py clean --file data/donnees_tourisme_france_exercice.csv")
    doc.add_paragraph()
    _code(doc, "# Générer le rapport complet avec les données propres")
    _code(doc, "python main.py report --file data/tourisme_brut.csv --output output/")
    doc.add_paragraph()
    _code(doc, "# Générer le document Word")
    _code(doc, "python generate_docx.py --output mon_rapport.docx")

    # ── Jupyter ───────────────────────────────────────────────────────────────
    _h2(doc, "7. Utiliser le notebook Jupyter")
    _para(doc,
        "Pour une analyse plus interactive avec les graphiques directement visibles "
        "dans le notebook, on peut ouvrir TP1.ipynb dans VS Code ou JupyterLab.")
    for b in [
        "Activer l'environnement virtuel (voir section 6).",
        "Ouvrir TP1.ipynb dans VS Code ou lancer : jupyter lab TP1.ipynb",
        "Exécuter toutes les cellules : Kernel → Restart & Run All",
    ]:
        _bullet(doc, b)
    _para(doc,
        "Le notebook fait exactement la même chose que la commande 'report', "
        "mais les graphiques s'affichent directement dans les cellules "
        "au lieu d'être sauvegardés en PNG.",
        space_before=4)


# ─────────────────────────────────────────────────────────────────────────────
# ASSEMBLAGE DU DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def build_document(output_path: str | Path,
                   charts_dir: Path = DEFAULT_CHARTS_DIR) -> None:
    doc = Document()

    # ── Mise en page A4 ──
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)

    # ── Style de base ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GREY

    # ── Page de couverture ──
    _cover(doc)

    # ── Table des matières ──
    p = _para(doc, "TABLE DES MATIÈRES", bold=True, size=16, color=NAVY)
    _hr(doc)
    doc.add_paragraph()
    _add_toc(doc)
    _para(doc,
          "⟵ Pour mettre à jour la table des matières dans Word : "
          "clic droit → « Mettre à jour les champs ».",
          italic=True, size=9, color=MID_GREY, space_before=10)
    _page_break(doc)

    # ── Parties ──
    _part1_rapport(doc, charts_dir)
    _part2_nettoyage(doc)
    _part3_doc(doc)

    # ── Sauvegarde ──
    output_path = Path(output_path)
    doc.save(output_path)
    print(f"\n✅  Document généré : {output_path.resolve()}")
    print(f"   Taille          : {output_path.stat().st_size / 1024:.1f} Ko")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Génère un document Word avec rapport + nettoyage + documentation."
    )
    parser.add_argument("--output", "-o", default="tourisme_rapport_complet.docx",
                        help="Nom du fichier DOCX (défaut : tourisme_rapport_complet.docx)")
    parser.add_argument("--charts", "-c", default="output",
                        help="Dossier contenant les graphiques PNG (défaut : output/)")
    args = parser.parse_args()
    build_document(args.output, charts_dir=Path(args.charts))


if __name__ == "__main__":
    main()
